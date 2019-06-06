#!/usr/bin/env python
# coding: utf-8

from docx import Document
import csv
import json
import os

with open("data.csv", "rb") as f:
    reader = csv.reader(f)
    header = reader.next()
fieldNames = header
identifier = header[0]


csvMetadata = open( 'data.csv', 'rU' )
ogDocument = Document('template.docx')

if not os.path.exists("results"):
    os.makedirs("results")


csvReader = csv.DictReader(
    csvMetadata,
    fieldnames = fieldNames
)

csvJSON = json.loads(
    json.dumps(
        [ row for row in csvReader ]
    )
)

del csvJSON[0]

for row in csvJSON:
    print "----"
    ogDocument.save("results/" + row[identifier] + '.docx')
    document = Document("results/" + row[identifier] + '.docx')
    print "Created document: results/" + row[identifier] + ".docx"
    for (key, value) in row.items():
        searchTerm = "$" + key
        for paragraph in document.paragraphs:
            if searchTerm in paragraph.text:
                print "Found: ", searchTerm, "and replacing it with:", value
                paragraph.text = paragraph.text.replace(searchTerm, value)
    document.save("results/" + row[identifier] + '.docx')
